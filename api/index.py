from supabase import create_client, Client
import httpx
import os
import csv
import base64
import difflib
import re
from io import BytesIO
from urllib.parse import urlencode
from flask import Flask, render_template, request, jsonify, session
from groq import Groq
from dotenv import load_dotenv


# Load secret environment API tokens locally
load_dotenv()

SUPABASE_URL = os.getenv("SUPABASE_URL")
SUPABASE_KEY = os.getenv("SUPABASE_ANON_KEY")
SUPABASE_SERVICE_KEY = os.getenv("SUPABASE_SERVICE_ROLE_KEY")

supabase_admin: Client | None = None
supabase: Client | None = None

if SUPABASE_URL and SUPABASE_SERVICE_KEY:
    supabase_admin = create_client(SUPABASE_URL, SUPABASE_SERVICE_KEY)

if SUPABASE_URL and SUPABASE_KEY:
    supabase = create_client(
        SUPABASE_URL,
        SUPABASE_KEY
)

# Robust absolute path calculation for Vercel Serverless environment and local VS Code testing
base_dir = os.path.dirname(os.path.abspath(__file__))
template_candidates = [
    os.path.abspath(os.path.join(base_dir, '..', 'templates')),
    os.path.abspath(os.path.join(base_dir, 'templates')),
    base_dir
]
template_dir = next(
    (path for path in template_candidates if os.path.exists(os.path.join(path, 'index.html'))),
    template_candidates[0]
)

app = Flask(__name__, template_folder=template_dir)
app.secret_key = os.environ.get("SECRET_KEY", "dev-secret-change-this-later")




def normalize_email(email):
    return (email or "").strip().lower()


# Supabase Auth helpers. This is the older working REST-based flow.
SUPABASE_REQUEST_TIMEOUT = 12.0

def supabase_is_configured():
    return bool(SUPABASE_URL and SUPABASE_KEY)

def supabase_headers(access_token=None):
    token = access_token or SUPABASE_KEY
    return {
        "apikey": SUPABASE_KEY or "",
        "Authorization": f"Bearer {token}",
        "Content-Type": "application/json"
    }

def supabase_auth_endpoint(path):
    return f"{str(SUPABASE_URL).rstrip('/')}/auth/v1/{path.lstrip('/')}"

def supabase_error_message(response, fallback="Supabase authentication request failed."):
    try:
        data = response.json()
    except Exception:
        return fallback
    return (
        data.get("msg")
        or data.get("message")
        or data.get("error_description")
        or data.get("error")
        or fallback
    )

def configured_site_origin():
    configured_url = (
        os.environ.get("SITE_URL")
        or os.environ.get("PUBLIC_SITE_URL")
        or os.environ.get("VERCEL_PROJECT_PRODUCTION_URL")
        or os.environ.get("VERCEL_URL")
        or ""
    ).strip()
    if configured_url:
        if not configured_url.startswith(("http://", "https://")):
            configured_url = f"https://{configured_url}"
        return configured_url.rstrip("/")
    return request.host_url.rstrip("/")

def store_authenticated_user(user, provider="email"):
    user = user or {}
    email = normalize_email(user.get("email"))
    user_id = user.get("id") or user.get("sub") or email
    session.clear()
    session["user_id"] = user_id
    session["user_email"] = email
    session["auth_provider"] = provider
    return {"id": user_id, "email": email, "provider": provider}



app.config['MAX_CONTENT_LENGTH'] = 20 * 1024 * 1024

# Initialize Groq Cloud Engine safely
api_key = os.environ.get("GROQ_API_KEY")
groq_client = Groq(api_key=api_key) if api_key else None


def extract_text_from_pdf(file_bytes):
    """Extract selectable text from a PDF file."""
    try:
        from pypdf import PdfReader

        reader = PdfReader(BytesIO(file_bytes))
        extracted_pages = []
        for page in reader.pages:
            extracted_pages.append(page.extract_text() or "")
        return "\n".join(extracted_pages).strip()
    except Exception:
        return ""


def extract_text_from_docx(file_bytes):
    """Extract paragraph and table text from a DOCX lab report."""
    try:
        from docx import Document

        document = Document(BytesIO(file_bytes))
        extracted_parts = []

        for paragraph in document.paragraphs:
            if paragraph.text.strip():
                extracted_parts.append(paragraph.text.strip())

        for table in document.tables:
            for row in table.rows:
                row_values = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_values.append(cell_text)
                if row_values:
                    extracted_parts.append(" | ".join(row_values))

        return "\n".join(extracted_parts).strip()
    except Exception:
        return ""


def extract_text_from_lab_image(image_bytes, mime_type="image/jpeg", context="lab report"):
    """Use Groq Vision to read lab report values from an uploaded image."""
    if not groq_client:
        return ""

    base64_encoded = base64.b64encode(image_bytes).decode('utf-8')

    vision_instruction = (
        f"Extract all medically relevant visible text from this {context}. Preserve test names, values, "
        "units, reference ranges, abnormal flags, dates, panel names, and radiology report wording if present. "
        "If this is a raw ultrasound image without report text, do not diagnose from the image. Only describe "
        "visible labels/text and state that radiology interpretation requires a clinician or official report."
    )

    chat_completion = groq_client.chat.completions.create(
        messages=[
            {
                "role": "user",
                "content": [
                    {"type": "text", "text": vision_instruction},
                    {
                        "type": "image_url",
                        "image_url": {
                            "url": f"data:{mime_type};base64,{base64_encoded}"
                        }
                    }
                ]
            }
        ],
        model="llama-3.2-11b-vision-preview",
        temperature=0.1
    )

    return chat_completion.choices[0].message.content.strip()


def extract_text_from_scanned_pdf(file_bytes, max_pages=5, context="lab report"):
    """Render PDF pages as images and use Groq Vision when selectable PDF text is unavailable."""
    try:
        import fitz

        document = fitz.open(stream=file_bytes, filetype="pdf")
        extracted_pages = []

        for page_index, page in enumerate(document):
            if page_index >= max_pages:
                extracted_pages.append(
                    f"[Only the first {max_pages} pages were image-read to keep processing within server limits.]"
                )
                break

            pixmap = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
            image_bytes = pixmap.tobytes("png")
            page_text = extract_text_from_lab_image(image_bytes, "image/png", context)
            if page_text:
                extracted_pages.append(f"--- Page {page_index + 1} ---\n{page_text}")

        return "\n\n".join(extracted_pages).strip()
    except Exception:
        return ""


def extract_uploaded_report_text(uploaded_file, context="lab report"):
    """Extract readable text from PDF, DOCX, PNG, JPG, or JPEG uploads."""
    if not uploaded_file or uploaded_file.filename == '':
        return "", ""

    filename = uploaded_file.filename
    filename_lower = filename.lower()
    file_bytes = uploaded_file.read()

    if filename_lower.endswith('.pdf'):
        extracted_text = extract_text_from_pdf(file_bytes)
        if len(extracted_text.strip()) < 50:
            scanned_text = extract_text_from_scanned_pdf(file_bytes, context=context)
            if scanned_text:
                extracted_text = scanned_text
        return extracted_text.strip(), filename

    if filename_lower.endswith('.docx'):
        return extract_text_from_docx(file_bytes), filename

    if filename_lower.endswith(('.png', '.jpg', '.jpeg')):
        mime_type = "image/png" if filename_lower.endswith('.png') else "image/jpeg"
        return extract_text_from_lab_image(file_bytes, mime_type, context), filename

    return "", filename


def extract_multiple_uploaded_reports(uploaded_files, context="liver-related report"):
    """Extract text from multiple uploaded liver-related files."""
    extracted_sections = []

    for index, uploaded_file in enumerate(uploaded_files, start=1):
        if not uploaded_file or not uploaded_file.filename:
            continue

        try:
            extracted_text, filename = extract_uploaded_report_text(uploaded_file, context=context)
            if extracted_text:
                extracted_sections.append(
                    f"===== Uploaded File {index}: {filename} =====\n{extracted_text}"
                )
            else:
                extracted_sections.append(
                    f"===== Uploaded File {index}: {filename} =====\n[No readable text could be extracted from this file.]"
                )
        except Exception as e:
            extracted_sections.append(
                f"===== Uploaded File {index}: {uploaded_file.filename} =====\n[File could not be processed: {str(e)}]"
            )

    return "\n\n".join(extracted_sections).strip()


def contains_negative_number(text):
    """Detect clearly negative numeric user inputs while allowing normal hyphenated labels/ranges."""
    if not text:
        return False

    compact_text = str(text).strip()
    if compact_text.startswith('-') and len(compact_text) > 1 and compact_text[1].isdigit():
        return True

    tokens = compact_text.replace('\n', ' ').split()
    for token in tokens:
        cleaned = token.strip(',;:()[]{}')
        if cleaned.startswith('-') and len(cleaned) > 1 and cleaned[1].isdigit():
            return True

    return False


def normalize_name(value):
    """Normalize compound names for matching."""
    return re.sub(r'[^a-z0-9]+', '', str(value).lower())


def get_dilirank_path():
    """Find the DILIrank CSV in common project locations."""
    candidates = [
        os.path.abspath(os.path.join(base_dir, '..', 'data', 'dilirank.csv')),
        os.path.abspath(os.path.join(base_dir, 'data', 'dilirank.csv')),
        os.path.abspath(os.path.join(os.getcwd(), 'data', 'dilirank.csv')),
        os.path.abspath(os.path.join(os.getcwd(), 'dilirank.csv'))
    ]

    for candidate in candidates:
        if os.path.exists(candidate):
            return candidate

    return None


def load_dilirank_records():
    """Load FDA DILIrank CSV, skipping the title row."""
    dilirank_path = get_dilirank_path()
    if not dilirank_path:
        return [], "DILIrank dataset was not found. Expected location: data/dilirank.csv"

    records = []

    try:
        with open(dilirank_path, "r", encoding="utf-8-sig", newline="") as file:
            reader = csv.reader(file)
            rows = list(reader)

        if len(rows) < 2:
            return [], "DILIrank dataset was found, but it does not contain enough rows."

        header = rows[1]
        data_rows = rows[2:]

        for row in data_rows:
            if not row or len(row) < 5:
                continue

            padded_row = row + [""] * (len(header) - len(row))
            row_dict = dict(zip(header, padded_row))

            compound_name = row_dict.get("CompoundName", "").strip()
            if not compound_name:
                continue

            records.append({
                "LTKBID": row_dict.get("LTKBID", "").strip(),
                "CompoundName": compound_name,
                "SeverityClass": row_dict.get("SeverityClass", "").strip(),
                "LabelSection": row_dict.get("LabelSection", "").strip(),
                "vDILIConcern": row_dict.get("vDILI-Concern", "").strip(),
                "Comment": row_dict.get("Comment", "").strip()
            })

        return records, ""
    except UnicodeDecodeError:
        try:
            with open(dilirank_path, "r", encoding="latin1", newline="") as file:
                reader = csv.reader(file)
                rows = list(reader)

            header = rows[1]
            data_rows = rows[2:]

            for row in data_rows:
                if not row or len(row) < 5:
                    continue

                padded_row = row + [""] * (len(header) - len(row))
                row_dict = dict(zip(header, padded_row))

                compound_name = row_dict.get("CompoundName", "").strip()
                if not compound_name:
                    continue

                records.append({
                    "LTKBID": row_dict.get("LTKBID", "").strip(),
                    "CompoundName": compound_name,
                    "SeverityClass": row_dict.get("SeverityClass", "").strip(),
                    "LabelSection": row_dict.get("LabelSection", "").strip(),
                    "vDILIConcern": row_dict.get("vDILI-Concern", "").strip(),
                    "Comment": row_dict.get("Comment", "").strip()
                })

            return records, ""
        except Exception as e:
            return [], f"Unable to read DILIrank dataset: {str(e)}"
    except Exception as e:
        return [], f"Unable to read DILIrank dataset: {str(e)}"


def concern_to_points(concern):
    """Convert DILI concern category into an educational baseline risk score."""
    value = str(concern).lower()

    if "most" in value:
        return 40
    if "ambiguous" in value:
        return 25
    if "less" in value:
        return 15
    if "no" in value:
        return 5

    return 10


def find_dilirank_match(query):
    """Find the best DILIrank match for a user-entered toxin/drug."""
    query = (query or "").strip()
    if not query:
        return {
            "found": False,
            "error": "",
            "message": "No toxin/drug name was provided.",
            "matches": [],
            "best_match": None
        }

    records, error = load_dilirank_records()
    if error:
        return {
            "found": False,
            "error": error,
            "message": error,
            "matches": [],
            "best_match": None
        }

    normalized_query = normalize_name(query)

    exact_matches = [
        record for record in records
        if normalize_name(record["CompoundName"]) == normalized_query
    ]

    if exact_matches:
        best = exact_matches[0]
        return {
            "found": True,
            "error": "",
            "message": "Exact DILIrank match found.",
            "matches": exact_matches[:5],
            "best_match": best
        }

    partial_matches = [
        record for record in records
        if normalized_query in normalize_name(record["CompoundName"]) or normalize_name(record["CompoundName"]) in normalized_query
    ]

    if partial_matches:
        best = partial_matches[0]
        return {
            "found": True,
            "error": "",
            "message": "Partial DILIrank match found.",
            "matches": partial_matches[:5],
            "best_match": best
        }

    names = [record["CompoundName"] for record in records]
    close_names = difflib.get_close_matches(query, names, n=5, cutoff=0.72)
    close_matches = [
        record for record in records
        if record["CompoundName"] in close_names
    ]

    if close_matches:
        best = close_matches[0]
        return {
            "found": True,
            "error": "",
            "message": "Close DILIrank match found. Verify that this is the intended toxin/drug.",
            "matches": close_matches[:5],
            "best_match": best
        }

    return {
        "found": False,
        "error": "",
        "message": "No close DILIrank match found for this toxin/drug.",
        "matches": [],
        "best_match": None
    }


def estimate_liver_risk_points(dili_match, amount, time_since_exposure, symptoms, typed_liver_labs, uploaded_report_text):
    """Create a simple educational risk estimate from DILI category plus user-supplied context."""
    points = 0
    reasons = []

    best_match = dili_match.get("best_match")
    if best_match:
        concern = best_match.get("vDILIConcern", "")
        dili_points = concern_to_points(concern)
        points += dili_points
        reasons.append(f"DILIrank category contributes {dili_points} points ({concern}).")
    else:
        points += 10
        reasons.append("No DILIrank match found, so dataset grounding is limited.")

    amount_text = (amount or "").lower()
    if not amount_text:
        points += 10
        reasons.append("Exposure amount is missing, which increases uncertainty.")
    elif any(word in amount_text for word in ["unknown", "large", "high", "overdose", "too much", "several", "many"]):
        points += 25
        reasons.append("Exposure amount suggests possible high-dose or uncertain-dose concern.")
    else:
        points += 5
        reasons.append("Exposure amount was provided, but exact clinical interpretation still requires context.")

    time_text = (time_since_exposure or "").lower()
    if not time_text:
        points += 5
        reasons.append("Time since exposure is missing.")
    else:
        points += 3
        reasons.append("Time since exposure was provided.")

    symptom_text = (symptoms or "").lower()
    severe_symptom_terms = [
        "jaundice", "yellow", "dark urine", "confusion", "severe abdominal", "right upper",
        "ruq", "bleeding", "vomiting blood", "faint", "unconscious"
    ]
    moderate_symptom_terms = [
        "nausea", "vomit", "fatigue", "itch", "itching", "abdominal pain", "loss of appetite",
        "pale stool", "clay stool"
    ]

    if any(term in symptom_text for term in severe_symptom_terms):
        points += 25
        reasons.append("Symptoms include liver red-flag terms such as jaundice, dark urine, confusion, or severe abdominal pain.")
    elif any(term in symptom_text for term in moderate_symptom_terms):
        points += 12
        reasons.append("Symptoms include possible liver-related warning signs.")
    elif symptom_text:
        points += 3
        reasons.append("Symptoms were provided, but no major liver red-flag terms were detected.")
    else:
        reasons.append("No symptoms were provided.")

    combined_report_text = f"{typed_liver_labs}\n{uploaded_report_text}".lower()
    severe_lab_terms = [
        "inr high", "inr elevated", "bilirubin high", "bilirubin elevated",
        "alt high", "ast high", "acute liver failure", "hepatic failure"
    ]

    if any(term in combined_report_text for term in severe_lab_terms):
        points += 20
        reasons.append("Typed/uploaded liver information includes concerning liver-lab wording.")
    elif combined_report_text.strip():
        points += 8
        reasons.append("Liver labs or report text were provided for interpretation.")
    else:
        reasons.append("No liver lab/report information was provided.")

    points = max(0, min(points, 100))

    if points <= 30:
        category = "Low educational concern"
    elif points <= 60:
        category = "Moderate educational concern"
    else:
        category = "High educational concern"

    return points, category, reasons


@app.errorhandler(404)
def not_found(error):
    return jsonify({
        "success": False,
        "error": "Route not found. Make sure your Flask server is running and you are using the correct /api/... URL."
    }), 404


@app.errorhandler(413)
def file_too_large(error):
    return jsonify({
        "success": False,
        "error": "Uploaded file is too large. Please upload a file under 20 MB total."
    }), 413


@app.errorhandler(415)
def unsupported_media_type(error):
    return jsonify({
        "success": False,
        "error": "Unsupported media type. This route expected JSON or form data, but received a different content type."
    }), 415


@app.errorhandler(500)
def internal_error(error):
    return jsonify({
        "success": False,
        "error": "Internal server error. Check your VS Code terminal for the full traceback."
    }), 500


@app.route('/')
def home():
    """Renders the main single-page interface application."""
    return render_template('index.html')


@app.route('/api/analyze', methods=['POST'])
def analyze_substance():
    """Handles Mode 1: Known Substance Risk Analyzer functionality."""
    if not groq_client:
        return jsonify({"success": False, "error": "Groq API key missing in environment backend config."})
    
    payload = request.get_json(silent=True) or {}
    substance = payload.get('substance', 'Unknown')
    duration = payload.get('duration', 'Unknown')
    contact_method = payload.get('contact_method', 'Unknown')

    system_prompt = (
        "You are an educational first-aid assistant specializing in household safety. "
        "Provide immediate non-medical guidance based on the given hazard parameters. "
        "Structure your response clearly with headers: 1) Immediate Action Needed, "
        "2) Common Symptoms/Side Effects to monitor over 24 hours, and 3) Risk Level Assessment (High/Low)."
    )
    
    user_message = f"Substance: {substance}\nExposure Duration: {duration}\nContact Method: {contact_method}"

    try:
        chat_completion = groq_client.chat.completions.create(
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_message}
            ],
            model="llama-3.1-8b-instant",
            temperature=0.2
        )
        ai_response = chat_completion.choices[0].message.content
        return jsonify({"success": True, "data": ai_response})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)})


@app.route('/api/identify', methods=['POST'])
def identify_symptoms():
    """Handles Mode 2: Symptom Identifier Mode functionality."""
    if not groq_client:
        return jsonify({"success": False, "error": "Groq API key missing in environment backend config."})
    
    payload = request.get_json(silent=True) or {}
    symptoms = payload.get('symptoms', 'None reported')
    context = payload.get('context', 'Unknown environment')

    system_prompt = (
        "You are an educational first-aid safety analyzer. Review the reported physical human symptoms "
        "and corresponding environment activity context. Suggest 2 or 3 common household chemicals, "
        "plants, or environmental hazards that match this description. Conclude by prompting the user "
        "to check safely if these specific items are present nearby. Keep it structural and clear."
    )
    
    user_message = f"Symptoms experienced: {symptoms}\nEnvironment Context: {context}"

    try:
        chat_completion = groq_client.chat.completions.create(
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_message}
            ],
            model="llama-3.1-8b-instant",
            temperature=0.2
        )
        ai_response = chat_completion.choices[0].message.content
        return jsonify({"success": True, "data": ai_response})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)})


@app.route('/api/scan', methods=['POST'])
def scan_label_image():
    """Handles Advanced Phase C: Vision Processing for chemical label photos."""
    if not groq_client:
        return jsonify({"success": False, "error": "Groq API engine client uninitialized."})
    
    if 'image' not in request.files:
        return jsonify({"success": False, "error": "No image file detected in form submission request."})
        
    uploaded_file = request.files['image']
    if uploaded_file.filename == '':
        return jsonify({"success": False, "error": "Selected filename is empty."})
        
    try:
        raw_bytes = uploaded_file.read()
        base64_encoded = base64.b64encode(raw_bytes).decode('utf-8')
        
        mime_type = "image/jpeg"
        if uploaded_file.filename.lower().endswith('.png'):
            mime_type = "image/png"
            
        vision_system_instruction = (
            "You are an emergency educational first-aid scanner. Examine the provided chemical label or ingredients text image. "
            "1) Identify the primary chemical compounds present. "
            "2) Highlight if any ingredients pose hazardous risks. "
            "3) Outline high-level non-medical first-aid parameters if skin, eye, or ingestion exposure occurs."
        )
        
        chat_completion = groq_client.chat.completions.create(
            messages=[
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": vision_system_instruction},
                        {
                            "type": "image_url",
                            "image_url": {
                                "url": f"data:{mime_type};base64,{base64_encoded}"
                            }
                        }
                    ]
                }
            ],
            model="llama-3.2-11b-vision-preview",
            temperature=0.15
        )
        
        vision_result = chat_completion.choices[0].message.content
        return jsonify({"success": True, "data": vision_result})
        
    except Exception as e:
        return jsonify({"success": False, "error": str(e)})


# CLINICAL LAB REPORT PATTERN ANALYZER ROUTE
@app.route('/api/lab_analyze', methods=['POST'])
def analyze_lab_reports():
    """Analyzes physiological panel values across core biological organ systems, including uploaded lab reports."""
    if not groq_client:
        return jsonify({"success": False, "error": "Groq API engine client uninitialized."})

    if request.is_json:
        payload = request.get_json(silent=True) or {}
        uploaded_file = None
    else:
        payload = request.form
        uploaded_file = request.files.get('lab_report') or request.files.get('file')
        
    toxin = payload.get('toxin', '').strip()
    amount = payload.get('amount', '').strip()
    lab_values = payload.get('lab_values', '').strip()

    uploaded_report_text = ""
    uploaded_filename = ""

    try:
        if uploaded_file and uploaded_file.filename:
            uploaded_report_text, uploaded_filename = extract_uploaded_report_text(uploaded_file, context="lab report")
    except Exception as e:
        return jsonify({"success": False, "error": f"Unable to read uploaded lab report: {str(e)}"})

    if not lab_values and not uploaded_report_text:
        return jsonify({
            "success": False,
            "error": "Please type lab values, upload a readable lab report, or provide both."
        })
    
    if contains_negative_number(lab_values) or contains_negative_number(amount):
        return jsonify({
            "success": True, 
            "data": "⚠️ CRITICAL EVALUATION ERROR\n\n[Confidence Score: 0%]\nReason: Invalid lab metrics provided. Lab results and exposure measurements cannot contain negative values. Please re-enter your official lab report statistics."
        })

    system_prompt = (
        "You are an educational physiological safety assistant. Analyze the user's submitted laboratory values "
        "and any uploaded lab report text against the suspected toxin exposure. Look for metabolic or stress "
        "patterns across major organ systems (such as Liver, Kidneys, Lungs, Heart, or Nervous System).\n\n"
        "STRICT COMPLIANCE RULES:\n"
        "1. DO NOT state definitively that the toxin caused these specific organ irregularities, as pre-existing diseases, "
        "chronic conditions, or baseline health histories are unverified.\n"
        "2. You must compute and explicitly output a 'Confidence Score' out of 100% at the very top of your response using these guidelines:\n"
        "   - If the toxin is completely unknown, fictional, or unrecognizable to toxicological literature, Confidence Score is 0%.\n"
        "   - If the user explicitly stated that the exposure amount is missing or if the amount field is completely unprovided/empty, reduce the confidence score significantly (assign no higher than 40-50%).\n"
        "   - If the lab values display typical physiological patterns correlated with a known exposure and all inputs are complete, assign a high confidence score (80-95%).\n"
        "3. Separate typed values from uploaded report findings when helpful, but combine them into one final interpretation.\n"
        "4. If the uploaded report is unclear or incomplete, state that limitation clearly.\n"
        "5. Provide a thorough analysis, then provide a concise summary at the end with the most important points for people who have come into contact with a very dangerous toxin. Provide a clear heading before adding the summary.\n"
        "6. Keep the response educational and encourage contacting poison control or emergency services for dangerous exposures."
    )
    
    user_message = (
        f"Suspected Toxin: {toxin if toxin else 'Not Provided'}\n"
        f"Exposure Amount: {amount if amount else 'NOT PROVIDED / MISSING'}\n"
        f"Typed Lab Metrics: {lab_values if lab_values else 'No typed lab metrics provided.'}\n"
        f"Uploaded Lab Report Filename: {uploaded_filename if uploaded_filename else 'No uploaded report provided.'}\n"
        f"Extracted Uploaded Lab Report Text:\n{uploaded_report_text if uploaded_report_text else 'No uploaded report text extracted.'}"
    )
    
    try:
        chat_completion = groq_client.chat.completions.create(
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_message}
            ],
            model="llama-3.1-8b-instant",
            temperature=0.2
        )
        ai_response = chat_completion.choices[0].message.content
        return jsonify({"success": True, "data": ai_response})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)})


# NEW FEATURE: LIVER-ONLY TOXIN IMPACT ANALYZER ROUTE
@app.route('/api/liver_impact', methods=['POST'])
def analyze_liver_impact():
    """Uses FDA DILIrank plus optional liver reports to estimate educational liver-risk patterns."""
    if not groq_client:
        return jsonify({"success": False, "error": "Groq API engine client uninitialized."})

    if request.is_json:
        payload = request.get_json(silent=True) or {}
        uploaded_files = []
    else:
        payload = request.form
        uploaded_files = request.files.getlist('liver_files')

    toxin = payload.get('toxin', '').strip()
    amount = payload.get('amount', '').strip()
    time_since_exposure = payload.get('time_since_exposure', '').strip()
    exposure_method = payload.get('exposure_method', '').strip()
    symptoms = payload.get('symptoms', '').strip()
    typed_liver_labs = payload.get('typed_liver_labs', '').strip()

    if not toxin:
        return jsonify({
            "success": False,
            "error": "Please enter a suspected toxin, medication, or chemical name."
        })

    if contains_negative_number(amount) or contains_negative_number(typed_liver_labs):
        return jsonify({
            "success": True,
            "data": "⚠️ INPUT CHECK FAILED\n\n[Educational Confidence Score: 0%]\nReason: Negative exposure or lab values were detected. Please re-check the entered liver-related values."
        })

    try:
        uploaded_report_text = extract_multiple_uploaded_reports(
            uploaded_files,
            context="liver blood test, liver panel, toxicology report, or liver ultrasound report"
        )
    except Exception as e:
        return jsonify({"success": False, "error": f"Unable to read uploaded liver report files: {str(e)}"})

    dili_match = find_dilirank_match(toxin)
    risk_points, risk_category, risk_reasons = estimate_liver_risk_points(
        dili_match,
        amount,
        time_since_exposure,
        symptoms,
        typed_liver_labs,
        uploaded_report_text
    )

    best_match = dili_match.get("best_match")
    matches = dili_match.get("matches", [])

    if best_match:
        dataset_summary = (
            f"DILIrank Lookup Status: {dili_match.get('message')}\n"
            f"Best Matched Compound: {best_match.get('CompoundName')}\n"
            f"LTKBID: {best_match.get('LTKBID')}\n"
            f"SeverityClass: {best_match.get('SeverityClass')}\n"
            f"FDA Label Section: {best_match.get('LabelSection')}\n"
            f"vDILI-Concern Category: {best_match.get('vDILIConcern')}\n"
            f"Dataset Comment: {best_match.get('Comment')}\n"
        )

        if len(matches) > 1:
            other_matches = "\n".join(
                [
                    f"- {match.get('CompoundName')} | {match.get('vDILIConcern')} | SeverityClass {match.get('SeverityClass')}"
                    for match in matches[1:5]
                ]
            )
            dataset_summary += f"Other Possible Dataset Matches:\n{other_matches}\n"
    else:
        dataset_summary = (
            f"DILIrank Lookup Status: {dili_match.get('message')}\n"
            "Best Matched Compound: None\n"
            "Dataset-grounded liver risk is limited because this toxin/drug was not found in DILIrank.\n"
        )

    risk_reason_text = "\n".join([f"- {reason}" for reason in risk_reasons])

    system_prompt = (
        "You are an educational liver-toxicology assistant for a liver-only toxin safety web app. "
        "You explain how a suspected toxin, medication, supplement, or chemical may affect the liver using "
        "FDA DILIrank dataset information, user exposure details, symptoms, and optional uploaded liver reports.\n\n"
        "VERY IMPORTANT SAFETY RULES:\n"
        "1. Do not diagnose the user.\n"
        "2. Do not claim that a future outcome is guaranteed.\n"
        "3. Use the phrase 'educational liver-risk forecast' instead of definitive prediction.\n"
        "4. If red-flag symptoms or severe lab patterns appear, tell the user to contact Poison Control, emergency services, or a clinician immediately.\n"
        "5. If a raw ultrasound image is uploaded, do not act like a radiologist. Only use official report text or visible labels.\n"
        "6. Keep the focus specifically on the liver, not every organ system.\n\n"
        "OUTPUT FORMAT:\n"
        "Start with:\n"
        "[Educational Liver-Risk Score: X/100]\n"
        "[Risk Category: Low / Moderate / High educational concern]\n"
        "[DILIrank Match: compound name or Not Found]\n\n"
        "Then include these sections:\n"
        "1. DILIrank Dataset Finding\n"
        "2. How This Toxin May Affect the Liver\n"
        "3. Expected Liver Lab Pattern\n"
        "4. User Liver Data / Uploaded Report Interpretation\n"
        "5. Educational Liver-Risk Forecast\n"
        "6. What Information Would Improve Confidence\n"
        "7. Short Summary of Most Important Points"
    )

    user_message = (
        f"User-entered suspected toxin/drug/chemical: {toxin}\n"
        f"Exposure amount/dose: {amount if amount else 'NOT PROVIDED / MISSING'}\n"
        f"Time since exposure: {time_since_exposure if time_since_exposure else 'NOT PROVIDED / MISSING'}\n"
        f"Exposure method: {exposure_method if exposure_method else 'NOT PROVIDED / MISSING'}\n"
        f"Symptoms: {symptoms if symptoms else 'No symptoms provided.'}\n"
        f"Typed liver labs: {typed_liver_labs if typed_liver_labs else 'No typed liver labs provided.'}\n\n"
        f"FDA DILIrank Dataset Information:\n{dataset_summary}\n"
        f"Educational Risk Score From App Formula: {risk_points}/100\n"
        f"Educational Risk Category From App Formula: {risk_category}\n"
        f"Risk Score Reasons:\n{risk_reason_text}\n\n"
        f"Extracted Text From Uploaded Liver-Related Files:\n{uploaded_report_text if uploaded_report_text else 'No liver-related files uploaded or no readable text extracted.'}"
    )

    try:
        chat_completion = groq_client.chat.completions.create(
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_message}
            ],
            model="llama-3.1-8b-instant",
            temperature=0.2
        )
        ai_response = chat_completion.choices[0].message.content
        return jsonify({"success": True, "data": ai_response})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)})


@app.route("/api/signup", methods=["POST"])
def signup():
    payload = request.get_json(silent=True) or {}
    email = normalize_email(payload.get("email"))
    password = payload.get("password") or ""

    if not supabase_is_configured():
        return jsonify({"success": False, "error": "Supabase is not configured. Add SUPABASE_URL and SUPABASE_ANON_KEY to your .env file."}), 500
    if not email or "@" not in email:
        return jsonify({"success": False, "error": "Please enter a valid email address."}), 400
    if len(password) < 8:
        return jsonify({"success": False, "error": "Password must be at least 8 characters long."}), 400

    try:
        response = httpx.post(
            supabase_auth_endpoint("signup"),
            headers=supabase_headers(),
            json={"email": email, "password": password},
            timeout=SUPABASE_REQUEST_TIMEOUT
        )
    except httpx.RequestError as e:
        return jsonify({"success": False, "error": f"Could not reach Supabase Auth: {str(e)}"}), 502

    if response.status_code >= 400:
        return jsonify({"success": False, "error": supabase_error_message(response, "Could not create this Supabase account.")}), response.status_code

    data = response.json()
    user = data.get("user") or {}
    session_data = data.get("session") or {}

    if session_data and session_data.get("access_token"):
        stored_user = store_authenticated_user(user, provider="email")
        return jsonify({"success": True, "logged_in": True, "message": "Account created and logged in successfully.", "user": {"email": stored_user.get("email")}})

    return jsonify({"success": True, "logged_in": False, "message": "Account created in Supabase. Check your email to confirm the account, then log in.", "user": {"email": user.get("email") or email}})


@app.route("/api/login", methods=["POST"])
def login():
    payload = request.get_json(silent=True) or {}
    email = normalize_email(payload.get("email"))
    password = payload.get("password") or ""

    if not supabase_is_configured():
        return jsonify({"success": False, "error": "Supabase is not configured. Add SUPABASE_URL and SUPABASE_ANON_KEY to your .env file."}), 500
    if not email or not password:
        return jsonify({"success": False, "error": "Email and password are required."}), 400

    try:
        response = httpx.post(
            supabase_auth_endpoint("token?grant_type=password"),
            headers=supabase_headers(),
            json={"email": email, "password": password},
            timeout=SUPABASE_REQUEST_TIMEOUT
        )
    except httpx.RequestError as e:
        return jsonify({"success": False, "error": f"Could not reach Supabase Auth: {str(e)}"}), 502

    if response.status_code >= 400:
        return jsonify({"success": False, "error": "Invalid email or password."}), 401

    data = response.json()
    user = data.get("user") or {}
    if not user:
        return jsonify({"success": False, "error": "Supabase did not return a user for this login."}), 502

    stored_user = store_authenticated_user(user, provider="email")
    return jsonify({"success": True, "message": "Logged in successfully.", "user": {"email": stored_user.get("email")}})


@app.route("/api/auth/google/start", methods=["GET"])
def google_login_start():
    if not supabase_is_configured():
        return jsonify({"success": False, "error": "Supabase is not configured. Add SUPABASE_URL and SUPABASE_ANON_KEY to your .env file."}), 500

    redirect_to = f"{configured_site_origin()}/api/auth/google/callback"
    query = urlencode({
        "provider": "google",
        "redirect_to": redirect_to
    })

    return jsonify({
        "success": True,
        "url": f"{supabase_auth_endpoint('authorize')}?{query}",
        "redirect_to": redirect_to
    })


# Backward-compatible route name, in case any older button still calls /api/google-login.
@app.route("/api/google-login", methods=["GET"])
def google_login():
    return google_login_start()


@app.route("/api/auth/google/callback", methods=["GET"])
def google_login_callback():
    return """<!DOCTYPE html>
<html lang=\"en\">
<head>
    <meta charset=\"UTF-8\">
    <meta name=\"viewport\" content=\"width=device-width, initial-scale=1.0\">
    <title>toxeye Google sign-in</title>
    <style>
        body { margin: 0; min-height: 100vh; display: flex; align-items: center; justify-content: center; background: #211a16; color: #f2ece5; font-family: Inter, ui-sans-serif, system-ui, -apple-system, BlinkMacSystemFont, \"Segoe UI\", sans-serif; }
        .card { width: min(24rem, calc(100% - 2rem)); border: 1.5px solid rgba(120, 94, 78, 0.65); border-radius: 1.25rem; padding: 1.5rem; background: rgba(44, 36, 31, 0.92); text-align: center; box-shadow: 0 18px 45px rgba(0, 0, 0, 0.22); }
        h1 { margin: 0 0 0.5rem 0; color: #caa982; font-size: 1.2rem; }
        p { margin: 0; color: #c9b8a8; font-size: 0.88rem; line-height: 1.5; }
    </style>
</head>
<body>
    <div class=\"card\">
        <h1>Finishing Google sign-in...</h1>
        <p id=\"status\">You can close this popup after sign-in completes.</p>
    </div>

    <script>
        async function finishGoogleLogin() {
            const statusEl = document.getElementById('status');
            const hashParams = new URLSearchParams(window.location.hash.slice(1));
            const queryParams = new URLSearchParams(window.location.search);

            const error = hashParams.get('error_description') || hashParams.get('error') || queryParams.get('error_description') || queryParams.get('error');
            if (error) {
                statusEl.textContent = error;
                if (window.opener) window.opener.postMessage({type: 'toxeye-google-login-result', success: false, error}, window.location.origin);
                return;
            }

            const accessToken = hashParams.get('access_token');
            const refreshToken = hashParams.get('refresh_token');

            if (!accessToken) {
                const missingTokenMessage = 'Google sign-in did not return an access token. Check the Supabase Google provider and redirect URL settings.';
                statusEl.textContent = missingTokenMessage;
                if (window.opener) window.opener.postMessage({type: 'toxeye-google-login-result', success: false, error: missingTokenMessage}, window.location.origin);
                return;
            }

            try {
                const response = await fetch('/api/auth/google/complete', {
                    method: 'POST',
                    headers: {'Content-Type': 'application/json'},
                    body: JSON.stringify({access_token: accessToken, refresh_token: refreshToken})
                });
                const result = await response.json();
                if (window.opener) window.opener.postMessage({type: 'toxeye-google-login-result', success: result.success, error: result.error || '', user: result.user || null}, window.location.origin);
                statusEl.textContent = result.success ? 'Google sign-in complete. Closing popup...' : (result.error || 'Google sign-in failed.');
                if (result.success) window.close();
            } catch (networkError) {
                const message = 'Could not finish Google sign-in.';
                statusEl.textContent = message;
                if (window.opener) window.opener.postMessage({type: 'toxeye-google-login-result', success: false, error: message}, window.location.origin);
            }
        }
        finishGoogleLogin();
    </script>
</body>
</html>"""


# Backward-compatible callback route. Supabase should use /api/auth/google/callback, but this keeps old settings from breaking.
@app.route("/auth/callback", methods=["GET"])
def auth_callback():
    return google_login_callback()


@app.route("/api/auth/google/complete", methods=["POST"])
def google_login_complete():
    payload = request.get_json(silent=True) or {}
    access_token = payload.get("access_token") or ""

    if not supabase_is_configured():
        return jsonify({"success": False, "error": "Supabase is not configured. Add SUPABASE_URL and SUPABASE_ANON_KEY to your .env file."}), 500
    if not access_token:
        return jsonify({"success": False, "error": "Missing Google access token from Supabase."}), 400

    try:
        response = httpx.get(
            supabase_auth_endpoint("user"),
            headers=supabase_headers(access_token),
            timeout=SUPABASE_REQUEST_TIMEOUT
        )
    except httpx.RequestError as e:
        return jsonify({"success": False, "error": f"Could not verify Google login with Supabase: {str(e)}"}), 502

    if response.status_code >= 400:
        return jsonify({"success": False, "error": supabase_error_message(response, "Could not verify this Google login with Supabase.")}), response.status_code

    user = response.json()
    stored_user = store_authenticated_user(user, provider="google")
    return jsonify({"success": True, "message": "Logged in with Google successfully.", "user": {"email": stored_user.get("email")}})


@app.route("/api/logout", methods=["POST"])
def logout():
    session.clear()
    return jsonify({"success": True, "message": "Logged out successfully."})


@app.route("/api/me", methods=["GET"])
def current_user():
    if "user_email" not in session:
        return jsonify({"success": True, "logged_in": False})

    return jsonify({
        "success": True,
        "logged_in": True,
        "user": {
            "email": session.get("user_email"),
            "provider": session.get("auth_provider", "email")
        }
    })

if __name__ == '__main__':
    app.run(debug=True)
